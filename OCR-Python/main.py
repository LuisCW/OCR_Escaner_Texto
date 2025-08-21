import os
import uuid
from typing import List
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import easyocr
import cv2
import numpy as np
from PIL import Image
from docx import Document
from docx.shared import Inches
import io

app = FastAPI(title="OCR API", description="API para OCR de texto manuscrito")

# Configurar CORS para permitir conexiones desde React Native
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # En producción, especificar dominios específicos
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Inicializar EasyOCR (soporta múltiples idiomas incluyendo español)
# EasyOCR es especialmente bueno para texto manuscrito
reader = easyocr.Reader(['es', 'en'])  # Español e Inglés

# Crear directorios necesarios
os.makedirs("uploads", exist_ok=True)
os.makedirs("output", exist_ok=True)

def preprocess_image(image_array):
    """
    Preprocesa la imagen para mejorar la precisión del OCR
    """
    # Convertir a escala de grises
    if len(image_array.shape) == 3:
        gray = cv2.cvtColor(image_array, cv2.COLOR_RGB2GRAY)
    else:
        gray = image_array
    
    # Aplicar filtro Gaussiano para reducir ruido
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)
    
    # Binarización adaptativa para mejorar el contraste
    binary = cv2.adaptiveThreshold(
        blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
    )
    
    # Operaciones morfológicas para limpiar la imagen
    kernel = np.ones((2, 2), np.uint8)
    cleaned = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)
    
    return cleaned

@app.get("/")
async def root():
    return {"message": "OCR API está funcionando correctamente"}

@app.post("/extract-text/")
async def extract_text(file: UploadFile = File(...)):
    """
    Extrae texto de una imagen usando OCR
    """
    try:
        # Verificar que el archivo sea una imagen
        if not file.content_type.startswith('image/'):
            raise HTTPException(status_code=400, detail="El archivo debe ser una imagen")
        
        # Leer la imagen
        contents = await file.read()
        image = Image.open(io.BytesIO(contents))
        
        # Convertir a array numpy
        image_array = np.array(image)
        
        # Preprocesar la imagen
        processed_image = preprocess_image(image_array)
        
        # Realizar OCR con EasyOCR
        results = reader.readtext(processed_image)
        
        # Extraer solo el texto
        extracted_text = []
        confidence_scores = []
        
        for (bbox, text, confidence) in results:
            if confidence > 0.3:  # Filtrar resultados con baja confianza
                extracted_text.append(text)
                confidence_scores.append(confidence)
        
        # Unir todo el texto
        full_text = " ".join(extracted_text)
        
        return {
            "text": full_text,
            "individual_texts": extracted_text,
            "confidence_scores": confidence_scores,
            "total_words": len(extracted_text)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error procesando la imagen: {str(e)}")

@app.post("/create-word-document/")
async def create_word_document(text: str, title: str = "Documento OCR"):
    """
    Crea un documento Word con el texto extraído
    """
    try:
        # Crear un nuevo documento
        doc = Document()
        
        # Agregar título
        doc.add_heading(title, 0)
        
        # Agregar información de metadatos
        doc.add_heading('Texto extraído por OCR', level=1)
        
        # Agregar el texto
        paragraph = doc.add_paragraph(text)
        
        # Generar nombre único para el archivo
        filename = f"documento_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join("output", filename)
        
        # Guardar el documento
        doc.save(filepath)
        
        return {
            "message": "Documento creado exitosamente",
            "filename": filename,
            "download_url": f"/download/{filename}"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creando el documento: {str(e)}")

@app.get("/download/{filename}")
async def download_file(filename: str):
    """
    Descarga un documento Word generado
    """
    filepath = os.path.join("output", filename)
    
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="Archivo no encontrado")
    
    return FileResponse(
        path=filepath,
        filename=filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.post("/process-image-to-word/")
async def process_image_to_word(
    file: UploadFile = File(...),
    title: str = "Documento OCR"
):
    """
    Procesa una imagen completa: extrae texto y crea documento Word
    """
    try:
        # Extraer texto
        contents = await file.read()
        image = Image.open(io.BytesIO(contents))
        image_array = np.array(image)
        processed_image = preprocess_image(image_array)
        
        results = reader.readtext(processed_image)
        
        extracted_text = []
        for (bbox, text, confidence) in results:
            if confidence > 0.3:
                extracted_text.append(text)
        
        full_text = " ".join(extracted_text)
        
        if not full_text.strip():
            raise HTTPException(status_code=400, detail="No se pudo extraer texto de la imagen")
        
        # Crear documento Word
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_heading('Texto extraído por OCR', level=1)
        doc.add_paragraph(full_text)
        
        # Agregar información adicional
        doc.add_paragraph(f"\nNúmero de palabras detectadas: {len(extracted_text)}")
        doc.add_paragraph(f"Procesado con EasyOCR")
        
        filename = f"documento_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join("output", filename)
        doc.save(filepath)
        
        return {
            "text": full_text,
            "word_count": len(extracted_text),
            "filename": filename,
            "download_url": f"/download/{filename}",
            "message": "Procesamiento completado exitosamente"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error en el procesamiento: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
